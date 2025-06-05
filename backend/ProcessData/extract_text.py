import pdfplumber
import docx
import pandas as pd
import tabula
import cv2
import numpy as np
from PIL import Image
import pytesseract
from dataclasses import dataclass
from typing import List, Dict, Any, Optional
import os
import re

@dataclass
class DocumentElement:
    type: str  # 'text', 'table', 'chart', 'formula'
    content: Any
    page_number: int
    metadata: Optional[Dict] = None

def extract_tables_from_pdf(pdf_path: str) -> List[DocumentElement]:
    """Extract tables from PDF using tabula-py"""
    tables = []
    try:
        # Read all tables from the PDF
        dfs = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)
        for page_num, df in enumerate(dfs, 1):
            if not df.empty:
                # Convert DataFrame to dictionary for better serialization
                table_dict = {
                    'data': df.to_dict(orient='records'),
                    'columns': df.columns.tolist(),
                    'shape': df.shape
                }
                tables.append(DocumentElement(
                    type='table',
                    content=table_dict,
                    page_number=page_num,
                    metadata={'table_index': len(tables)}
                ))
    except Exception as e:
        print(f"Error extracting tables: {str(e)}")
    return tables

def ensure_charts_dir():
    """Tạo thư mục lưu trữ biểu đồ nếu chưa tồn tại"""
    charts_dir = "./extracted_charts"
    if not os.path.exists(charts_dir):
        os.makedirs(charts_dir)
        print(f"Đã tạo thư mục {charts_dir}")
    return charts_dir

def extract_charts_from_pdf(pdf_path: str) -> List[DocumentElement]:
    """Extract charts from PDF using OpenCV and Tesseract"""
    charts = []
    try:
        # Tạo thư mục lưu trữ biểu đồ
        charts_dir = ensure_charts_dir()
        
        with pdfplumber.open(pdf_path) as pdf:
            # Lấy tên file PDF không có phần mở rộng
            pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
            
            for page_num, page in enumerate(pdf.pages, 1):
                # Convert page to image
                img = page.to_image()
                img_array = np.array(img.original)
                
                # Convert to grayscale
                gray = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
                
                # Apply threshold to get binary image
                _, binary = cv2.threshold(gray, 240, 255, cv2.THRESH_BINARY_INV)
                
                # Find contours
                contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                # Get page dimensions
                height, width = img_array.shape[:2]
                page_area = height * width
                
                for idx, contour in enumerate(contours):
                    # Calculate contour properties
                    area = cv2.contourArea(contour)
                    x, y, w, h = cv2.boundingRect(contour)
                    aspect_ratio = float(w)/h if h > 0 else 0
                    
                    # Filter contours based on multiple criteria
                    if (area > 10000 and  # Minimum area
                        area < page_area * 0.8 and  # Maximum area (80% of page)
                        0.2 < aspect_ratio < 5 and  # Reasonable aspect ratio
                        w > 100 and h > 100):  # Minimum dimensions
                        
                        # Extract the region
                        chart_img = img_array[y:y+h, x:x+w]
                        
                        # Check if the region contains enough non-white pixels
                        non_white_pixels = np.sum(chart_img < 240)
                        if non_white_pixels > (w * h * 0.1):  # At least 10% non-white pixels
                            
                            # Tạo tên file ảnh với định dạng: pdf_name_page_chart.png
                            chart_filename = f"{pdf_name}_page{page_num}_chart{idx}.png"
                            chart_path = os.path.join(charts_dir, chart_filename)
                            
                            # Save chart image
                            cv2.imwrite(chart_path, chart_img)
                            
                            # Extract text from chart using OCR
                            try:
                                chart_text = pytesseract.image_to_string(chart_img, lang='eng+vie')
                            except:
                                chart_text = "Không thể trích xuất text từ biểu đồ"
                            
                            # Add to charts list
                            charts.append(DocumentElement(
                                type='chart',
                                content={
                                    'image_path': chart_path,
                                    'text': chart_text,
                                    'position': {'x': x, 'y': y, 'width': w, 'height': h},
                                    'properties': {
                                        'area': area,
                                        'aspect_ratio': aspect_ratio,
                                        'non_white_ratio': non_white_pixels / (w * h)
                                    }
                                },
                                page_number=page_num,
                                metadata={'chart_index': idx}
                            ))
    except Exception as e:
        print(f"Error extracting charts: {str(e)}")
    return charts

def extract_formulas_from_text(text: str, page_num: int) -> List[DocumentElement]:
    """Extract mathematical formulas from text using enhanced regex patterns"""
    formulas = []
    
    # Enhanced patterns for mathematical expressions
    patterns = [
        # LaTeX style formulas
        r'\$([^$]+)\$',  # Inline math
        r'\$\$([^$]+)\$\$',  # Display math
        # Equations and expressions
        r'([a-zA-Z][a-zA-Z0-9]*\s*=\s*[^=]+)',  # Equations with variables
        r'([a-zA-Z][a-zA-Z0-9]*\s*\([^)]+\)\s*=\s*[^=]+)',  # Function definitions
        r'([a-zA-Z][a-zA-Z0-9]*\s*\([^)]+\)\s*[+\-*/]\s*[^=]+)',  # Function expressions
        # Fractions and ratios
        r'(\d+/\d+)',  # Simple fractions
        r'([a-zA-Z0-9]+\s*/\s*[a-zA-Z0-9]+)',  # Fractions with variables
        # Square roots and radicals
        r'√\s*\(([^)]+)\)',  # Square root
        r'sqrt\s*\(([^)]+)\)',  # Square root (alternative notation)
        # Powers and exponents
        r'([a-zA-Z0-9]+\s*\^\s*[a-zA-Z0-9]+)',  # Power with ^
        r'([a-zA-Z0-9]+\s*\*\*\s*[a-zA-Z0-9]+)',  # Power with **
        # Subscripts
        r'([a-zA-Z0-9]+\s*_\s*[a-zA-Z0-9]+)',  # Subscript
        # Summations and products
        r'(∑|∏)\s*([^∑∏]+)',  # Summation or product
        # Integrals
        r'∫\s*([^∫]+)',  # Integral
        # Greek letters and special symbols
        r'([α-ωΑ-Ω][a-zA-Z0-9]*)',  # Greek letters
        # Matrices
        r'\[([^\[\]]+)\]',  # Matrix notation
        # Inequalities
        r'([a-zA-Z0-9]+\s*[<>≤≥]\s*[a-zA-Z0-9]+)',  # Inequalities
        # Percentages
        r'(\d+%)',  # Percentages
        # Units and measurements
        r'(\d+\s*(?:m|kg|s|A|K|mol|cd|Hz|N|Pa|J|W|C|V|Ω|F|H|T|Wb|lm|lx|Bq|Gy|Sv|kat|L|t|ha|eV|u|Da|bar|atm|mmHg|Torr|psi|cal|kcal|Wh|kWh|dB|ppm|ppb|ppt|mol/L|g/L|mg/L|μg/L|ng/L|pg/L|mol/m³|g/m³|mg/m³|μg/m³|ng/m³|pg/m³|mol/kg|g/kg|mg/kg|μg/kg|ng/kg|pg/kg|mol/mol|g/g|mg/g|μg/g|ng/g|pg/g|mol/m²|g/m²|mg/m²|μg/m²|ng/m²|pg/m²|mol/s|g/s|mg/s|μg/s|ng/s|pg/s|mol/min|g/min|mg/min|μg/min|ng/min|pg/min|mol/h|g/h|mg/h|μg/h|ng/h|pg/h|mol/d|g/d|mg/d|μg/d|ng/d|pg/d|mol/wk|g/wk|mg/wk|μg/wk|ng/wk|pg/wk|mol/mo|g/mo|mg/mo|μg/mo|ng/mo|pg/mo|mol/yr|g/yr|mg/yr|μg/yr|ng/yr|pg/yr))',  # Units
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            formula = match.group(1).strip()
            # Get context (text before and after the formula)
            start = max(0, match.start() - 150)  # Increased context window
            end = min(len(text), match.end() + 150)
            context = text[start:end]
            
            # Enhanced formula type identification
            formula_type = "unknown"
            if "=" in formula:
                formula_type = "equation"
            elif "/" in formula:
                formula_type = "fraction"
            elif "^" in formula or "**" in formula:
                formula_type = "power"
            elif "_" in formula:
                formula_type = "subscript"
            elif "√" in formula or "sqrt" in formula:
                formula_type = "square_root"
            elif "∑" in formula or "∏" in formula:
                formula_type = "summation_or_product"
            elif "∫" in formula:
                formula_type = "integral"
            elif any(greek in formula for greek in "αβγδεζηθικλμνξοπρστυφχψως"):
                formula_type = "greek_notation"
            elif "[" in formula and "]" in formula:
                formula_type = "matrix"
            elif any(op in formula for op in "<>≤≥"):
                formula_type = "inequality"
            elif "%" in formula:
                formula_type = "percentage"
            elif re.search(r'\d+\s*(?:m|kg|s|A|K|mol|cd|Hz|N|Pa|J|W|C|V|Ω|F|H|T|Wb|lm|lx|Bq|Gy|Sv|kat|L|t|ha|eV|u|Da|bar|atm|mmHg|Torr|psi|cal|kcal|Wh|kWh|dB|ppm|ppb|ppt|mol/L|g/L|mg/L|μg/L|ng/L|pg/L|mol/m³|g/m³|mg/m³|μg/m³|ng/m³|pg/m³|mol/kg|g/kg|mg/kg|μg/kg|ng/kg|pg/kg|mol/mol|g/g|mg/g|μg/g|ng/g|pg/g|mol/m²|g/m²|mg/m²|μg/m²|ng/m²|pg/m²|mol/s|g/s|mg/s|μg/s|ng/s|pg/s|mol/min|g/min|mg/min|μg/min|ng/min|pg/min|mol/h|g/h|mg/h|μg/h|ng/h|pg/h|mol/d|g/d|mg/d|μg/d|ng/d|pg/d|mol/wk|g/wk|mg/wk|μg/wk|ng/wk|pg/wk|mol/mo|g/mo|mg/mo|μg/mo|ng/mo|pg/mo|mol/yr|g/yr|mg/yr|μg/yr|ng/yr|pg/yr)', formula):
                formula_type = "measurement"
            
            # Extract variables and constants
            variables = set(re.findall(r'[a-zA-Z][a-zA-Z0-9]*', formula))
            constants = set(re.findall(r'\b\d+(?:\.\d+)?\b', formula))
            
            # Create LaTeX-like representation using string replacement
            latex = formula
            replacements = [
                ('^', '^{'),
                ('_', '_{'),
                ('√', r'\sqrt{'),
                ('sqrt', r'\sqrt{'),
                ('∑', r'\sum'),
                ('∏', r'\prod'),
                ('∫', r'\int'),
                ('α', r'\alpha'),
                ('β', r'\beta'),
                ('γ', r'\gamma'),
                ('δ', r'\delta'),
                ('ε', r'\varepsilon'),
                ('ζ', r'\zeta'),
                ('η', r'\eta'),
                ('θ', r'\theta'),
                ('ι', r'\iota'),
                ('κ', r'\kappa'),
                ('λ', r'\lambda'),
                ('μ', r'\mu'),
                ('ν', r'\nu'),
                ('ξ', r'\xi'),
                ('ο', r'\omicron'),
                ('π', r'\pi'),
                ('ρ', r'\rho'),
                ('σ', r'\sigma'),
                ('τ', r'\tau'),
                ('υ', r'\upsilon'),
                ('φ', r'\phi'),
                ('χ', r'\chi'),
                ('ψ', r'\psi'),
                ('ω', r'\omega'),
                ('≤', r'\leq'),
                ('≥', r'\geq'),
                ('≠', r'\neq'),
                ('±', r'\pm'),
                ('∞', r'\infty'),
                ('∂', r'\partial'),
                ('∇', r'\nabla'),
                ('∅', r'\emptyset'),
                ('∈', r'\in'),
                ('∉', r'\notin'),
                ('⊂', r'\subset'),
                ('⊃', r'\supset'),
                ('∪', r'\cup'),
                ('∩', r'\cap'),
                ('∅', r'\emptyset'),
                ('∀', r'\forall'),
                ('∃', r'\exists'),
                ('∄', r'\nexists'),
                ('∝', r'\propto'),
                ('∞', r'\infty'),
                ('ℵ', r'\aleph'),
                ('ℜ', r'\Re'),
                ('ℑ', r'\Im'),
                ('℘', r'\wp'),
                ('ℵ', r'\aleph'),
                ('ℶ', r'\beth'),
                ('ℷ', r'\gimel'),
                ('ℸ', r'\daleth')
            ]
            
            for old, new in replacements:
                latex = latex.replace(old, new)
            
            formulas.append(DocumentElement(
                type='formula',
                content={
                    'formula': formula,
                    'context': context,
                    'latex': latex,
                    'variables': list(variables),
                    'constants': list(constants)
                },
                page_number=page_num,
                metadata={
                    'formula_index': len(formulas),
                    'formula_type': formula_type,
                    'variables': list(variables),
                    'constants': list(constants),
                    'has_equals': '=' in formula,
                    'has_fraction': '/' in formula,
                    'has_power': '^' in formula or '**' in formula,
                    'has_subscript': '_' in formula,
                    'has_square_root': '√' in formula or 'sqrt' in formula,
                    'has_summation': '∑' in formula,
                    'has_product': '∏' in formula,
                    'has_integral': '∫' in formula,
                    'has_greek': any(greek in formula for greek in "αβγδεζηθικλμνξοπρστυφχψως"),
                    'has_matrix': '[' in formula and ']' in formula,
                    'has_inequality': any(op in formula for op in "<>≤≥"),
                    'has_percentage': '%' in formula,
                    'has_measurement': bool(re.search(r'\d+\s*(?:m|kg|s|A|K|mol|cd|Hz|N|Pa|J|W|C|V|Ω|F|H|T|Wb|lm|lx|Bq|Gy|Sv|kat|L|t|ha|eV|u|Da|bar|atm|mmHg|Torr|psi|cal|kcal|Wh|kWh|dB|ppm|ppb|ppt|mol/L|g/L|mg/L|μg/L|ng/L|pg/L|mol/m³|g/m³|mg/m³|μg/m³|ng/m³|pg/m³|mol/kg|g/kg|mg/kg|μg/kg|ng/kg|pg/kg|mol/mol|g/g|mg/g|μg/g|ng/g|pg/g|mol/m²|g/m²|mg/m²|μg/m²|ng/m²|pg/m²|mol/s|g/s|mg/s|μg/s|ng/s|pg/s|mol/min|g/min|mg/min|μg/min|ng/min|pg/min|mol/h|g/h|mg/h|μg/h|ng/h|pg/h|mol/d|g/d|mg/d|μg/d|ng/d|pg/d|mol/wk|g/wk|mg/wk|μg/wk|ng/wk|pg/wk|mol/mo|g/mo|mg/mo|μg/mo|ng/mo|pg/mo|mol/yr|g/yr|mg/yr|μg/yr|ng/yr|pg/yr)', formula))
                }
            ))
    
    return formulas

def extract_from_pdf(pdf_path: str) -> List[DocumentElement]:
    """Extract all elements (text, tables, charts, formulas) from PDF"""
    elements = []
    
    # Extract text and formulas
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                # Extract formulas from text
                formulas = extract_formulas_from_text(text, page_num)
                elements.extend(formulas)
                
                # Add text element
                elements.append(DocumentElement(
                    type='text',
                    content=text,
                    page_number=page_num
                ))
    
    # Extract tables
    tables = extract_tables_from_pdf(pdf_path)
    elements.extend(tables)
    
    # Extract charts
    charts = extract_charts_from_pdf(pdf_path)
    elements.extend(charts)
    
    # Sort elements by page number
    elements.sort(key=lambda x: x.page_number)
    return elements

def extract_from_docx(path: str) -> List[DocumentElement]:
    """Extract elements from DOCX file"""
    elements = []
    doc = docx.Document(path)
    
    for page_num, paragraph in enumerate(doc.paragraphs, 1):
        if paragraph.text.strip():
            elements.append(DocumentElement(
                type='text',
                content=paragraph.text,
                page_number=page_num
            ))
    
    # Extract tables from DOCX
    for table_num, table in enumerate(doc.tables, 1):
        table_data = []
        for row in table.rows:
            table_data.append([cell.text for cell in row.cells])
        
        if table_data:
            elements.append(DocumentElement(
                type='table',
                content={
                    'data': table_data,
                    'shape': (len(table_data), len(table_data[0]) if table_data else 0)
                },
                page_number=table_num,
                metadata={'table_index': table_num}
            ))
    
    return elements
